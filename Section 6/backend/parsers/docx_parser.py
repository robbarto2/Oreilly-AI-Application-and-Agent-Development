from docx import Document
from backend.parsers.base import DocumentParser, ParsedContent
from typing import List
import logging

logger = logging.getLogger(__name__)


class DOCXParser(DocumentParser):
    """Parser for Word (DOCX) files preserving heading hierarchy"""

    def supports(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))

    async def parse(self, file_path: str) -> ParsedContent:
        """
        Parse DOCX file preserving heading hierarchy.

        Extracts:
        - Document structure with heading levels
        - Paragraphs and lists
        - Tables
        - Distinguishes proposals from published courses
        """
        try:
            doc = Document(file_path)
            sections = self._build_hierarchy(doc)

            # Detect document type from content
            doc_type = self._detect_document_type(doc)

            root = ParsedContent(
                title=self._extract_title(doc) or f"Document from {file_path}",
                content=f"Word document with {len(doc.paragraphs)} paragraphs",
                item_type="document",
                children=sections,
                metadata={
                    "total_paragraphs": len(doc.paragraphs),
                    "total_tables": len(doc.tables),
                    "document_type": doc_type,
                },
                sequence_order=0,
            )

            logger.info(f"Parsed DOCX: {len(sections)} sections from {file_path}")
            return root

        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {e}")
            raise

    def _extract_title(self, doc: Document) -> str:
        """Extract document title from first heading or paragraph"""
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading 1') or para.style.name == 'Title':
                return para.text.strip()
            if para.text.strip():
                return para.text.strip()
        return ""

    def _detect_document_type(self, doc: Document) -> str:
        """Detect if this is a proposal, published course, or experimental notes"""
        text_lower = " ".join(p.text.lower() for p in doc.paragraphs[:10])

        if any(keyword in text_lower for keyword in ['proposal', 'future', 'roadmap', 'planned']):
            return "proposal"
        elif any(keyword in text_lower for keyword in ['course', 'syllabus', 'curriculum', 'learning objectives']):
            return "published_course"
        elif any(keyword in text_lower for keyword in ['experiment', 'draft', 'notes', 'ideas']):
            return "experimental"
        else:
            return "unknown"

    def _build_hierarchy(self, doc: Document) -> List[ParsedContent]:
        """Build hierarchical structure from headings and paragraphs"""
        sections = []
        current_section = None
        current_subsection = None
        para_order = 0

        for para in doc.paragraphs:
            para_order += 1
            style = para.style.name
            text = para.text.strip()

            if not text:
                continue

            if style.startswith('Heading 1') or style == 'Title':
                # New top-level section
                if current_section:
                    sections.append(current_section)

                current_section = ParsedContent(
                    title=text,
                    content="",
                    item_type="section",
                    children=[],
                    metadata={"heading_level": 1},
                    sequence_order=para_order,
                )
                current_subsection = None

            elif style.startswith('Heading 2'):
                # Subsection
                if current_section:
                    subsection = ParsedContent(
                        title=text,
                        content="",
                        item_type="section",
                        children=[],
                        metadata={"heading_level": 2},
                        sequence_order=para_order,
                    )
                    current_section.children.append(subsection)
                    current_subsection = subsection

            elif style.startswith('Heading'):
                # Lower-level heading
                level = int(style[-1]) if style[-1].isdigit() else 3
                subsection = ParsedContent(
                    title=text,
                    content="",
                    item_type="section",
                    children=[],
                    metadata={"heading_level": level},
                    sequence_order=para_order,
                )

                if current_subsection:
                    current_subsection.children.append(subsection)
                elif current_section:
                    current_section.children.append(subsection)

            else:
                # Regular paragraph content
                para_node = ParsedContent(
                    title="",
                    content=text,
                    item_type="paragraph",
                    metadata={"style": style},
                    sequence_order=para_order,
                )

                if current_subsection:
                    current_subsection.children.append(para_node)
                elif current_section:
                    current_section.children.append(para_node)
                else:
                    # No section yet, create a default one
                    if not current_section:
                        current_section = ParsedContent(
                            title="Introduction",
                            content="",
                            item_type="section",
                            children=[],
                            metadata={"heading_level": 1},
                            sequence_order=0,
                        )
                    current_section.children.append(para_node)

        # Add final section
        if current_section:
            sections.append(current_section)

        # Process tables
        for idx, table in enumerate(doc.tables):
            table_content = self._extract_table(table)
            table_node = ParsedContent(
                title=f"Table {idx + 1}",
                content=table_content,
                item_type="table",
                metadata={"rows": len(table.rows), "cols": len(table.columns)},
                sequence_order=para_order + idx + 1,
            )
            sections.append(table_node)

        return sections

    def _extract_table(self, table) -> str:
        """Extract table data as formatted text"""
        table_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_text.append(row_text)
        return "\n".join(table_text)
